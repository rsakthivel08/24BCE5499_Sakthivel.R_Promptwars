"""
app/extraction/document_parser.py
───────────────────────────────────
Extracts raw text from uploaded PDF and DOCX files.
Uses PyMuPDF (fitz) for PDFs and python-docx for Word documents.
"""
from __future__ import annotations

from pathlib import Path

from app.utils.logging_config import get_logger

logger = get_logger(__name__)


def extract_text_from_pdf(file_path: Path) -> str:
    """Extract all text from a PDF using pdfplumber."""
    try:
        import pdfplumber

        pages_text: list[str] = []

        with pdfplumber.open(str(file_path)) as pdf:
            for page_num, page in enumerate(pdf.pages, start=1):
                text = page.extract_text()

                if text and text.strip():
                    pages_text.append(
                        f"--- Page {page_num} ---\n{text.strip()}"
                    )

        result = "\n".join(pages_text)

        if not result.strip():
            raise ValueError("No extractable text found in the PDF.")

        logger.info(
            "pdf_extracted",
            path=str(file_path),
            pages=len(pages_text)
        )

        return result

    except Exception as exc:
        logger.error(
            "pdf_extraction_failed",
            path=str(file_path),
            error=str(exc)
        )
        raise


def extract_text_from_docx(file_path: Path) -> str:
    """Extract all text from a DOCX file using python-docx."""
    try:
        from docx import Document

        doc = Document(str(file_path))
        paragraphs: list[str] = [para.text for para in doc.paragraphs if para.text.strip()]
        # Also extract table cells
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    paragraphs.append(row_text)
        result = "\n".join(paragraphs)
        logger.info("docx_extracted", path=str(file_path), paragraphs=len(paragraphs))
        return result
    except Exception as exc:
        logger.error("docx_extraction_failed", path=str(file_path), error=str(exc))
        raise


def extract_text(file_path: Path) -> str:
    """
    Auto-detect file type and extract text.
    Supports: .pdf, .docx, .doc, .txt
    """
    suffix = file_path.suffix.lower()

    if suffix == ".pdf":
        return extract_text_from_pdf(file_path)
    elif suffix in (".docx", ".doc"):
        return extract_text_from_docx(file_path)
    elif suffix == ".txt":
        return file_path.read_text(encoding="utf-8", errors="replace")
    else:
        raise ValueError(f"Unsupported file type: {suffix}. Supported: pdf, docx, txt")
